from datetime import datetime
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document

from services.auth import require_login
from services.repository import (
    EXPORTS_DIR,
    build_material_summary,
    get_materials,
    get_profile,
    get_programs,
    get_tasks,
    init_db,
)
from services.ui import setup_page


setup_page("导出中心")
require_login()
init_db()
st.title("导出中心")

EXCEL_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


PROGRAM_COLUMNS = {
    "id": "ID",
    "region": "国家/地区",
    "school_name": "学校名称",
    "program_name": "项目名称",
    "degree_type": "学位类型",
    "academic_direction": "专业方向",
    "website": "项目官网链接",
    "source_url": "信息来源链接",
    "verified_date": "核验日期",
    "is_shared": "可见范围",
    "tuition": "学费",
    "living_cost": "生活费预估",
    "application_fee": "申请费",
    "deadline": "截止日期",
    "gpa_requirement": "GPA 要求",
    "toefl_requirement": "托福要求",
    "gre_gmat_requirement": "GRE/GMAT 要求",
    "recommendation_requirement": "推荐信要求",
    "essay_requirement": "文书要求",
    "scholarship_available": "是否有奖学金",
    "employment_notes": "就业导向备注",
    "category": "申请分类",
    "status": "当前状态",
    "notes": "备注",
}

TASK_COLUMNS = {
    "id": "ID",
    "task_name": "任务名称",
    "related_program": "所属学校/项目",
    "task_type": "任务类型",
    "deadline": "截止日期",
    "priority": "优先级",
    "status": "当前状态",
    "notes": "备注",
}

PROFILE_LABELS = [
    ("目标入学季", "target_intake"),
    ("姓名", "name"),
    ("本科专业", "undergraduate_major"),
    ("年级", "grade"),
    ("当前 GPA", "current_gpa"),
    ("GPA 满分制", "gpa_scale"),
    ("托福当前分数", "toefl_current"),
    ("托福目标分数", "toefl_target"),
    ("GRE/GMAT 状态", "gre_gmat_status"),
    ("目标国家/地区", "target_regions"),
    ("目标专业方向", "target_majors"),
    ("预算范围", "budget_range"),
    ("实习经历", "internships"),
    ("科研/论文经历", "research_experience"),
    ("比赛/项目经历", "competitions_projects"),
    ("家教/兼职经历", "tutoring_part_time"),
    ("股票投资经历", "stock_investment"),
    ("职业目标", "career_goals"),
    ("备注", "notes"),
]


def excel_bytes(df: pd.DataFrame, sheet_name: str) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return output.getvalue()


def docx_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def format_programs(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame(columns=PROGRAM_COLUMNS.values())
    output = df.copy()
    output["scholarship_available"] = output["scholarship_available"].map(lambda value: "是" if int(value or 0) else "否")
    output["is_shared"] = output["is_shared"].map(lambda value: "公共" if int(value or 0) else "仅自己")
    return output[list(PROGRAM_COLUMNS.keys())].rename(columns=PROGRAM_COLUMNS)


def format_tasks(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame(columns=TASK_COLUMNS.values())
    return df[list(TASK_COLUMNS.keys())].rename(columns=TASK_COLUMNS)


def build_profile_doc(profile: dict) -> Document:
    doc = Document()
    doc.add_heading("个人申请档案", level=0)
    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, key in PROFILE_LABELS:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(profile.get(key) or "")
    return doc


def build_materials_doc(materials: pd.DataFrame) -> Document:
    doc = Document()
    doc.add_heading("文书素材库", level=0)
    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if materials.empty:
        doc.add_paragraph("暂无文书素材。")
        return doc
    for row in materials.to_dict("records"):
        doc.add_heading(row.get("experience_name") or "未命名经历", level=1)
        for paragraph in build_material_summary(row).split("\n\n"):
            doc.add_paragraph(paragraph)
    return doc


program_export = format_programs(get_programs())
task_export = format_tasks(get_tasks())
profile_doc = build_profile_doc(get_profile())
materials_doc = build_materials_doc(get_materials())

program_bytes = excel_bytes(program_export, "选校表")
task_bytes = excel_bytes(task_export, "申请时间线")
profile_bytes = docx_bytes(profile_doc)
materials_bytes = docx_bytes(materials_doc)

col1, col2 = st.columns(2)
with col1:
    st.download_button(
        "下载选校表 Excel",
        data=program_bytes,
        file_name="StudyAbroadOS_选校表.xlsx",
        mime=EXCEL_MIME,
        use_container_width=True,
    )
    st.download_button(
        "下载文书素材 Word",
        data=materials_bytes,
        file_name="StudyAbroadOS_文书素材库.docx",
        mime=DOCX_MIME,
        use_container_width=True,
    )

with col2:
    st.download_button(
        "下载申请时间线 Excel",
        data=task_bytes,
        file_name="StudyAbroadOS_申请时间线.xlsx",
        mime=EXCEL_MIME,
        use_container_width=True,
    )
    st.download_button(
        "下载个人申请档案 Word",
        data=profile_bytes,
        file_name="StudyAbroadOS_个人申请档案.docx",
        mime=DOCX_MIME,
        use_container_width=True,
    )

st.divider()

if st.button("生成并保存全部文件到 exports 文件夹", type="primary"):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    files = {
        EXPORTS_DIR / f"StudyAbroadOS_选校表_{timestamp}.xlsx": program_bytes,
        EXPORTS_DIR / f"StudyAbroadOS_申请时间线_{timestamp}.xlsx": task_bytes,
        EXPORTS_DIR / f"StudyAbroadOS_文书素材库_{timestamp}.docx": materials_bytes,
        EXPORTS_DIR / f"StudyAbroadOS_个人申请档案_{timestamp}.docx": profile_bytes,
    }
    for path, content in files.items():
        path.write_bytes(content)
    st.success("已保存到 exports 文件夹。")
    for path in files:
        st.code(str(path))

st.caption("导出的学校、项目要求和费用均来自你在系统中手动输入的数据。")
